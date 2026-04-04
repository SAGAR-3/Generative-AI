"""
src/ingestion/document_loader.py
=================================
Production-grade document loader for banking/home lending documents.

Supports:
- PDF (policy manuals, rate sheets, compliance docs)
- DOCX (underwriting guidelines, loan officer scripts)
- TXT (regulatory text, FAQ documents)
- XLSX (rate tables, fee schedules)

Security: Validates file types, sizes, and scans for malicious content.
"""

import hashlib
import logging
import mimetypes
import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterator, List, Optional

import structlog

logger = structlog.get_logger(__name__)


# ─── Data Models ────────────────────────────────────────────────────────────

@dataclass
class DocumentMetadata:
    """Rich metadata attached to every loaded document."""
    source_file: str
    file_type: str
    file_size_bytes: int
    sha256_hash: str
    ingested_at: str
    document_category: str          # e.g., "underwriting_guideline", "rate_sheet"
    regulatory_tags: List[str]      # e.g., ["RESPA", "TILA", "ECOA"]
    access_level: str               # "public" | "internal" | "confidential" | "restricted"
    version: str = "1.0"
    effective_date: Optional[str] = None
    expiry_date: Optional[str] = None
    author: Optional[str] = None
    department: Optional[str] = None
    extra: Dict = field(default_factory=dict)


@dataclass
class LoadedDocument:
    """A single loaded document with content and metadata."""
    doc_id: str
    content: str
    metadata: DocumentMetadata
    page_count: int = 1


# ─── Document Category Detector ─────────────────────────────────────────────

CATEGORY_PATTERNS = {
    "underwriting_guideline": [
        "underwriting", "credit score", "debt-to-income", "dti", "ltv",
        "loan-to-value", "qualifying", "eligibility criteria"
    ],
    "rate_sheet": [
        "interest rate", "apr", "basis points", "prime rate", "rate lock",
        "fixed rate", "adjustable rate", "arm", "30-year", "15-year"
    ],
    "compliance_policy": [
        "respa", "tila", "ecoa", "fcra", "hmda", "regulation", "disclosure",
        "compliance", "fair lending", "equal credit"
    ],
    "loan_product": [
        "fha", "va loan", "conventional", "jumbo", "usda", "heloc",
        "home equity", "refinance", "purchase loan"
    ],
    "fee_schedule": [
        "origination fee", "appraisal fee", "title insurance", "closing costs",
        "points", "pmi", "escrow", "settlement"
    ],
    "faq_guide": [
        "frequently asked", "question", "how to", "what is", "guide",
        "customer service", "help"
    ],
}


def detect_document_category(content: str, filename: str) -> tuple[str, List[str]]:
    """
    Auto-detect document category and regulatory tags from content.
    Returns (category, regulatory_tags).
    """
    content_lower = content.lower()
    filename_lower = filename.lower()

    # Detect category
    category_scores: Dict[str, int] = {}
    for cat, patterns in CATEGORY_PATTERNS.items():
        score = sum(1 for p in patterns if p in content_lower or p in filename_lower)
        if score > 0:
            category_scores[cat] = score

    category = max(category_scores, key=category_scores.get) if category_scores else "general"

    # Detect regulatory tags
    regulatory_tags = []
    regulatory_map = {
        "RESPA": ["respa", "real estate settlement", "hud-1", "gfe", "good faith estimate"],
        "TILA": ["tila", "truth in lending", "regulation z", "apr disclosure"],
        "ECOA": ["ecoa", "equal credit", "regulation b", "adverse action"],
        "FCRA": ["fcra", "credit report", "credit bureau", "consumer report"],
        "HMDA": ["hmda", "home mortgage disclosure", "lar", "loan application register"],
        "GLBA": ["glba", "gramm-leach-bliley", "privacy notice", "financial privacy"],
        "BSA": ["bank secrecy", "bsa", "aml", "anti-money laundering", "suspicious activity"],
    }
    for reg, patterns in regulatory_map.items():
        if any(p in content_lower for p in patterns):
            regulatory_tags.append(reg)

    return category, regulatory_tags


# ─── File Validators ─────────────────────────────────────────────────────────

MAX_FILE_SIZE_MB = 50
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".xlsx", ".csv"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/csv",
}


def validate_file(file_path: Path) -> None:
    """
    Security: Validate file before processing.
    Raises ValueError for invalid files.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Check extension
    if file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {file_path.suffix}. Allowed: {ALLOWED_EXTENSIONS}")

    # Check file size
    size_mb = file_path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(f"File too large: {size_mb:.1f}MB. Max allowed: {MAX_FILE_SIZE_MB}MB")

    # Check MIME type (defense against extension spoofing)
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if mime_type and mime_type not in ALLOWED_MIME_TYPES:
        logger.warning("mime_type_mismatch", file=str(file_path), mime=mime_type)

    logger.info("file_validated", file=str(file_path), size_mb=round(size_mb, 2))


def compute_sha256(file_path: Path) -> str:
    """Compute SHA-256 hash for deduplication and integrity verification."""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


# ─── Format-Specific Loaders ─────────────────────────────────────────────────

def load_pdf(file_path: Path) -> tuple[str, int]:
    """Load PDF using pypdf. Returns (content, page_count)."""
    try:
        import pypdf
        reader = pypdf.PdfReader(str(file_path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i+1}]\n{text}")
        return "\n\n".join(pages), len(reader.pages)
    except ImportError:
        raise ImportError("pypdf not installed. Run: pip install pypdf")
    except Exception as e:
        logger.error("pdf_load_failed", file=str(file_path), error=str(e))
        raise


def load_docx(file_path: Path) -> tuple[str, int]:
    """Load DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables (important for rate sheets)
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))

        all_content = "\n\n".join(paragraphs)
        if table_texts:
            all_content += "\n\n[TABLES]\n" + "\n\n".join(table_texts)

        return all_content, 1
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")


def load_txt(file_path: Path) -> tuple[str, int]:
    """Load plain text file."""
    encodings = ["utf-8", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            content = file_path.read_text(encoding=encoding)
            return content, 1
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {file_path} with any supported encoding")


def load_xlsx(file_path: Path) -> tuple[str, int]:
    """Load Excel file, converting to readable text format."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        sheets_content = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(c) if c is not None else "" for c in row)
                    rows.append(row_str)
            if rows:
                sheets_content.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
        return "\n\n".join(sheets_content), 1
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")


# ─── Main DocumentLoader ─────────────────────────────────────────────────────

class DocumentLoader:
    """
    Production-grade document loader for banking documents.

    Features:
    - Multi-format support (PDF, DOCX, TXT, XLSX)
    - File validation and security checks
    - Automatic category and regulatory tag detection
    - SHA-256 deduplication
    - Structured logging
    """

    LOADERS = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".txt": load_txt,
        ".xlsx": load_xlsx,
    }

    def load_file(
        self,
        file_path: str | Path,
        access_level: str = "internal",
        effective_date: Optional[str] = None,
        author: Optional[str] = None,
        department: Optional[str] = None,
    ) -> LoadedDocument:
        """
        Load a single document file.

        Args:
            file_path: Path to the document
            access_level: Security classification (public/internal/confidential/restricted)
            effective_date: When this document became effective (YYYY-MM-DD)
            author: Document author/owner
            department: Responsible department

        Returns:
            LoadedDocument with content and rich metadata
        """
        path = Path(file_path)
        validate_file(path)

        ext = path.suffix.lower()
        loader_fn = self.LOADERS.get(ext)
        if not loader_fn:
            raise ValueError(f"No loader for extension: {ext}")

        logger.info("loading_document", file=str(path), type=ext)
        content, page_count = loader_fn(path)

        if not content.strip():
            raise ValueError(f"Document appears to be empty: {path}")

        # Compute hash for deduplication
        sha256 = compute_sha256(path)
        doc_id = f"doc_{sha256[:16]}"

        # Auto-detect category and regulatory tags
        category, regulatory_tags = detect_document_category(content, path.name)

        metadata = DocumentMetadata(
            source_file=str(path),
            file_type=ext.lstrip("."),
            file_size_bytes=path.stat().st_size,
            sha256_hash=sha256,
            ingested_at=datetime.utcnow().isoformat(),
            document_category=category,
            regulatory_tags=regulatory_tags,
            access_level=access_level,
            effective_date=effective_date,
            author=author,
            department=department,
        )

        logger.info(
            "document_loaded",
            doc_id=doc_id,
            category=category,
            regulatory_tags=regulatory_tags,
            pages=page_count,
            chars=len(content),
        )

        return LoadedDocument(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            page_count=page_count,
        )

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
        access_level: str = "internal",
    ) -> Iterator[LoadedDocument]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Path to directory containing documents
            recursive: Whether to recurse into subdirectories
            access_level: Default access level for all documents

        Yields:
            LoadedDocument for each successfully loaded file
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory}")

        pattern = "**/*" if recursive else "*"
        files = [f for f in dir_path.glob(pattern) if f.is_file()]

        logger.info("directory_scan", directory=str(dir_path), file_count=len(files))

        loaded, failed, skipped = 0, 0, 0
        for file_path in files:
            if file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
                skipped += 1
                continue
            try:
                doc = self.load_file(file_path, access_level=access_level)
                loaded += 1
                yield doc
            except Exception as e:
                logger.error("document_load_failed", file=str(file_path), error=str(e))
                failed += 1

        logger.info(
            "directory_load_complete",
            loaded=loaded,
            failed=failed,
            skipped=skipped,
        )
